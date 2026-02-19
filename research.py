import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from collections import Counter
import re
import matplotlib as mpl
from wordcloud import WordCloud
from io import BytesIO
from docx import Document

# --- 1. ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤‡∏Å‡∏≤‡∏£‡πÅ‡∏™‡∏î‡∏á‡∏ú‡∏•‡∏†‡∏≤‡∏©‡∏≤‡πÑ‡∏ó‡∏¢ ---
font_path = "Kanit-Regular.ttf" 

def setup_font():
    try:
        mpl.font_manager.fontManager.addfont(font_path)
        prop = mpl.font_manager.FontProperties(fname=font_path)
        mpl.rc('font', family=prop.get_name())
        mpl.rcParams['axes.unicode_minus'] = False 
        return True
    except:
        return False

# --- 2. ‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏≠‡∏≤‡∏£‡∏°‡∏ì‡πå‡πÅ‡∏ö‡∏ö‡∏ú‡∏™‡∏°‡∏ú‡∏™‡∏≤‡∏ô ---
def analyze_sentiment_thai(text):
    pos_words = ['‡∏î‡∏µ', '‡πÄ‡∏´‡πá‡∏ô‡∏î‡πâ‡∏ß‡∏¢', '‡∏†‡∏π‡∏°‡∏¥‡πÉ‡∏à', '‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏™‡∏∏‡∏Ç', '‡∏û‡∏±‡∏í‡∏ô‡∏≤', '‡∏õ‡∏£‡∏∞‡πÇ‡∏¢‡∏ä‡∏ô‡πå', '‡∏¢‡∏±‡πà‡∏á‡∏¢‡∏∑‡∏ô', '‡∏û‡∏≠‡πÄ‡∏û‡∏µ‡∏¢‡∏á', '‡∏õ‡∏£‡∏∞‡∏´‡∏¢‡∏±‡∏î']
    neg_words = ['‡πÑ‡∏°‡πà‡∏î‡∏µ', '‡∏õ‡∏±‡∏ç‡∏´‡∏≤', '‡πÅ‡∏¢‡πà', '‡∏¢‡∏≤‡∏Å‡∏•‡∏≥‡∏ö‡∏≤‡∏Å', '‡∏Ç‡∏≤‡∏î‡πÅ‡∏Ñ‡∏•‡∏ô', '‡∏≠‡∏∏‡∏õ‡∏™‡∏£‡∏£‡∏Ñ', '‡∏´‡∏ô‡∏µ‡πâ‡∏™‡∏¥‡∏ô', '‡πÄ‡∏î‡∏∑‡∏≠‡∏î‡∏£‡πâ‡∏≠‡∏ô', '‡πÄ‡∏™‡∏µ‡∏¢‡∏î‡∏≤‡∏¢']
    pos_score = sum(1 for w in pos_words if w in text)
    neg_score = sum(1 for w in neg_words if w in text)
    if pos_score > neg_score: return "‡∏ö‡∏ß‡∏Å (Positive) üòä"
    elif neg_score > pos_score: return "‡∏•‡∏ö (Negative) üòü"
    else: return "‡∏õ‡∏Å‡∏ï‡∏¥ / ‡πÄ‡∏õ‡πá‡∏ô‡∏Å‡∏•‡∏≤‡∏á üòê"

# --- 3. ‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô MS Word ---
def create_word_report(filename, sentiment, summary, keywords_df, original_text):
    doc = Document()
    doc.add_heading(f'‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏ß‡∏¥‡∏à‡∏±‡∏¢: {filename}', 0)
    
    doc.add_heading('1. ‡∏ú‡∏•‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏≠‡∏≤‡∏£‡∏°‡∏ì‡πå', level=1)
    doc.add_paragraph(sentiment)
    
    doc.add_heading('2. ‡∏™‡∏£‡∏∏‡∏õ‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç', level=1)
    for s in summary:
        doc.add_paragraph(s, style='List Bullet')
        
    doc.add_heading('3. ‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç‡∏ó‡∏µ‡πà‡∏û‡∏ö (Top Keywords)', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç (‡∏¢‡∏≤‡∏ß >= 5, ‡∏ã‡πâ‡∏≥ >= 3)'
    hdr_cells[1].text = '‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏Ñ‡∏£‡∏±‡πâ‡∏á'
    for index, row in keywords_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['‡∏Ñ‡∏≥'])
        row_cells[1].text = str(row['‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏Ñ‡∏£‡∏±‡πâ‡∏á'])

    doc.add_heading('4. ‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ï‡πâ‡∏ô‡∏â‡∏ö‡∏±‡∏ö', level=1)
    doc.add_paragraph(original_text[:1000] + "..." if len(original_text) > 1000 else original_text)
        
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 4. ‡∏Å‡∏≤‡∏£‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤ Library ---
try:
    from pythainlp.summarize import summarize
    from pythainlp.tokenize import word_tokenize
    from pythainlp.corpus import thai_stopwords
    THAI_READY = True
except ImportError:
    THAI_READY = False

st.set_page_config(layout="wide", page_title="Full Research Analysis Tool")
st.title("üìÇ ‡∏£‡∏∞‡∏ö‡∏ö‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏á‡∏≤‡∏ô‡∏ß‡∏¥‡∏à‡∏±‡∏¢‡∏â‡∏ö‡∏±‡∏ö‡∏™‡∏°‡∏ö‡∏π‡∏£‡∏ì‡πå (All-in-One)")

if not THAI_READY:
    st.error("‚ùå ‡∏û‡∏ö‡∏Ç‡πâ‡∏≠‡∏ú‡∏¥‡∏î‡∏û‡∏•‡∏≤‡∏î‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡πÇ‡∏´‡∏•‡∏î Library")
    st.stop()

setup_font()
uploaded_files = st.file_uploader("‡∏≠‡∏±‡∏õ‡πÇ‡∏´‡∏•‡∏î‡πÑ‡∏ü‡∏•‡πå‡∏ö‡∏ó‡∏™‡∏±‡∏°‡∏†‡∏≤‡∏©‡∏ì‡πå (.txt)", type=['txt'], accept_multiple_files=True)

if uploaded_files:
    comparison_list = [] # ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏ß‡∏°
    
    for file in uploaded_files:
        text = file.read().decode("utf-8")
        tokens = word_tokenize(text, keep_whitespace=False)
        stop_words = list(thai_stopwords())
        extra_stop = ['‡πÄ‡∏ô‡∏≤‡∏∞', '‡∏ô‡∏∞', '‡∏Ñ‡∏£‡∏±‡∏ö', '‡∏Ñ‡πà‡∏∞', '‡∏≠‡∏∑‡∏°', '‡πÄ‡∏≠‡πà‡∏≠']
        stop_words.extend(extra_stop)
        
        # ‡∏Å‡∏£‡∏≠‡∏á‡∏Ñ‡∏≥: ‡∏¢‡∏≤‡∏ß >= 5 ‡πÅ‡∏•‡∏∞‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πà‡∏™‡∏±‡∏ç‡∏•‡∏±‡∏Å‡∏©‡∏ì‡πå
        filtered_by_length = [t.strip() for t in tokens if t.strip() and t not in stop_words and len(t.strip()) >= 5 and not re.match(r'^[0-9\W]+$', t)]
        word_counts = Counter(filtered_by_length)
        # ‡∏Å‡∏£‡∏≠‡∏á‡∏ã‡πâ‡∏≥ >= 3 ‡∏Ñ‡∏£‡∏±‡πâ‡∏á
        filtered_final = [word for word in filtered_by_length if word_counts[word] >= 3]
        
        s_label = analyze_sentiment_thai(text)
        try:
            brief = summarize(text, n=2)
        except:
            brief = ["‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏™‡∏£‡∏∏‡∏õ‡πÑ‡∏î‡πâ"]

        # ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÄ‡∏Ç‡πâ‡∏≤‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡∏£‡∏ß‡∏°
        comparison_list.append({
            "‡∏ä‡∏∑‡πà‡∏≠‡πÑ‡∏ü‡∏•‡πå": file.name,
            "‡πÇ‡∏ó‡∏ô‡∏Ñ‡∏ß‡∏≤‡∏°‡∏£‡∏π‡πâ‡∏™‡∏∂‡∏Å": s_label,
            "‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç‡∏´‡∏•‡∏±‡∏Å": Counter(filtered_final).most_common(1)[0][0] if filtered_final else "‡πÑ‡∏°‡πà‡∏û‡∏ö"
        })

        with st.expander(f"üìë ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î: {file.name}", expanded=True):
            tab1, tab2 = st.tabs(["üìä ‡∏™‡∏ñ‡∏¥‡∏ï‡∏¥‡πÅ‡∏•‡∏∞ Word Cloud", "üìÑ ‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ï‡πâ‡∏ô‡∏â‡∏ö‡∏±‡∏ö"])
            
            with tab1:
                c1, c2 = st.columns(2)
                with c1:
                    st.subheader("üí° ‡∏Å‡∏≤‡∏£‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤")
                    st.write(f"**‡∏Ñ‡∏ß‡∏≤‡∏°‡∏£‡∏π‡πâ‡∏™‡∏∂‡∏Å:** {s_label}")
                    st.write("**‡∏™‡∏£‡∏∏‡∏õ‡∏õ‡∏£‡∏∞‡πÄ‡∏î‡πá‡∏ô:**")
                    for b in brief: st.write(f"üìå {b}")
                    
                    if filtered_final:
                        wc = WordCloud(width=800, height=400, background_color="white", regexp=r"[\u0e00-\u0e7f]+", font_path=font_path).generate(" ".join(filtered_final))
                        fig, ax = plt.subplots()
                        ax.imshow(wc, interpolation='bilinear')
                        ax.axis("off")
                        st.pyplot(fig)
                        
                        buf = BytesIO()
                        fig.savefig(buf, format="png")
                        st.download_button(label="üíæ ‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏π‡∏õ Word Cloud (PNG)", data=buf.getvalue(), file_name=f"cloud_{file.name}.png", mime="image/png")
                
                with c2:
                    st.subheader("üìà ‡∏™‡∏ñ‡∏¥‡∏ï‡∏¥‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç")
                    final_counts = Counter(filtered_final).most_common(12)
                    df_counts = pd.DataFrame(final_counts, columns=['‡∏Ñ‡∏≥', '‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏Ñ‡∏£‡∏±‡πâ‡∏á'])
                    st.table(df_counts)
                    
                    word_report = create_word_report(file.name, s_label, brief, df_counts, text)
                    st.download_button(label="üìÑ ‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏™‡∏£‡∏∏‡∏õ (MS Word)", data=word_report, file_name=f"report_{file.name}.docx")

            with tab2:
                st.subheader("‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ï‡πâ‡∏ô‡∏â‡∏ö‡∏±‡∏ö")
                st.info("‡πÅ‡∏™‡∏î‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ö‡∏≤‡∏á‡∏™‡πà‡∏ß‡∏ô‡∏à‡∏≤‡∏Å‡πÑ‡∏ü‡∏•‡πå‡∏Ç‡∏≠‡∏á‡∏Ñ‡∏∏‡∏ì")
                st.text_area(label="Content Viewer", value=text, height=300)

    # --- ‡∏™‡πà‡∏ß‡∏ô‡∏™‡∏∏‡∏î‡∏ó‡πâ‡∏≤‡∏¢: ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡∏£‡∏ß‡∏° (Cross-Case Table) ---
    st.divider()
    st.subheader("üìã ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡∏ó‡∏∏‡∏Å‡πÄ‡∏Ñ‡∏™ (Sentiment Analysis Summary)")
    st.table(pd.DataFrame(comparison_list))
else:
    st.info("üëã ‡∏Å‡∏£‡∏∏‡∏ì‡∏≤‡∏≠‡∏±‡∏õ‡πÇ‡∏´‡∏•‡∏î‡πÑ‡∏ü‡∏•‡πå‡∏ö‡∏ó‡∏™‡∏±‡∏°‡∏†‡∏≤‡∏©‡∏ì‡πå‡πÄ‡∏û‡∏∑‡πà‡∏≠‡πÄ‡∏£‡∏¥‡πà‡∏°‡∏Å‡∏≤‡∏£‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå")
